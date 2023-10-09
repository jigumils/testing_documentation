#import File
import streamlit as st
import pandas as pd
from string import ascii_uppercase as abjad
from datetime import datetime
import io
from databases import user

##CPTFX,CPTPS,RPTDAILY
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH as wap, WD_TAB_LEADER, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm,Pt,Inches

##TestScript, Deflog, Summary Testing
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.datavalidation import DataValidation

##Dashboard
import matplotlib.pyplot as plt
import seaborn as sb

"""
Function List
format:
    #Function shortdesc
    #menu: A,B,...
    #function declaration
"""

def buat_space(doc,n):
  p = doc.add_paragraph()
  p.paragraph_format.line_spacing = n
  p.paragraph_format.space_after = 0

def buat_empty_table(doc, border_style=True):
  table = doc.add_table(rows=1, cols=1)
  if border_style:
    table.style = 'Table Grid'

def buat_cptfx():
  doc = docx.Document()
  for i in range(16):
    list_data = ["Project Name", "Module Name", "Defect ID", "Description", "Date", "Tester"]
    table = doc.add_table(rows=len(list_data), cols=2)
    table.style = 'Table Grid'

    for num,data in enumerate(list_data):
      row = table.rows[num].cells
      row[0].text = data
      if num==0:
        row[1].text = input_data["Project Name"]
      elif num==2:
        row[1].text = f"Defect-{i+1}"
      elif num==4:
        bulan = ["Januari", "Februari", "Maret",
                 "April", "Mei", "Juni",
                 "Juli", "Agustus", "September",
                 "Oktober", "November", "Desember"]
        row[1].text = f'{bulan[datetime.now().month-1]} {datetime.now().year}'
      elif num==5:
        row[1].text = df_user[df_user["NIP"] == input_data["Tester"]]["Nickname"][0]

    #buat_space(doc,1)
    buat_empty_table(doc)
    buat_space(doc,2)
  return doc



#Initialization, Directories, and Databases
##Initialization
user_dict = user.user_dict
df_user = pd.DataFrame(user_dict)

#Input Program
input_data = {"Project Name":str,
              "Project Version": str,
              "List Module":[],
              "Tester":int,
              "Testing Fase (SIT/UAT)":str}


#Main Program
st.write('Test Apps')

input_data["Tester"] = int(st.text_input(label="NIP TESTER", value=user_dict["NIP"][0]))
input_data["Project Name"] = "Project " + st.text_input(label="Project Name", placeholder="QRIS Acquirer")
input_data["Project Version"] = "v" + st.text_input(label="Project Version", placeholder = "1.2")
input_data["List Module"] = st.text_input(label="List Module", placeholder="Module_1,Module_2,Module_3").split(",")
input_data["Testing Fase (SIT/UAT)"] = st.radio(label="Testing Fase (SIT/UAT)", options=["SIT","UAT"])

st.write(input_data)

doc_download = buat_cptfx()

bio = io.BytesIO()
doc_download.save(bio)
if doc_download:
    st.download_button(
        label="Click here to download",
        data=bio.getvalue(),
        file_name=f'28-CPTESTFX-{input_data["Project Name"]}-{input_data["Project Version"]}.docx',
        mime="docx"
    )

#Output Program
